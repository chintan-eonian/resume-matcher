"""
Convert test resumes from TXT to PDF and DOCX formats
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import os

# Try to import font for PDF (DejaVu or fallback)
try:
    FONT_PATH = None
    # Check for DejaVu font (commonly available)
    possible_fonts = [
        "C:/Windows/Fonts/DejaVuSansCondensed.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/times.ttf",
    ]
    for font_path in possible_fonts:
        if os.path.exists(font_path):
            FONT_PATH = font_path
            break
except:
    FONT_PATH = None

def convert_txt_to_docx(txt_file, output_file):
    """Convert a TXT resume to DOCX format."""
    try:
        # Read TXT content
        with open(txt_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create new DOCX document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Process content line by line
        lines = content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add spacing
                if current_paragraph:
                    doc.add_paragraph()
                current_paragraph = None
                continue
            
            # Check if it's a heading (all caps or followed by dashes/equals)
            is_heading = (
                line.isupper() and len(line) > 5 and ' ' in line
            ) or line.endswith((':', '-', '=')) or (
                len(line) < 50 and line.isupper()
            )
            
            if is_heading:
                # Add heading
                para = doc.add_heading(line.rstrip(':-=').strip(), level=2)
                para.style.font.size = Pt(14)
                para.style.font.bold = True
                current_paragraph = None
            elif line.startswith(('Email:', 'Phone:', 'LinkedIn:')) or '@' in line:
                # Contact info - bold first part
                if ':' in line:
                    parts = line.split(':', 1)
                    para = doc.add_paragraph()
                    run1 = para.add_run(parts[0] + ':')
                    run1.bold = True
                    if len(parts) > 1:
                        para.add_run(parts[1])
                else:
                    para = doc.add_paragraph(line)
                current_paragraph = para
            elif line.startswith('•') or line.startswith('-'):
                # Bullet point
                para = doc.add_paragraph(line.lstrip('•- '), style='List Bullet')
                current_paragraph = para
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                current_paragraph = para
        
        # Save document
        doc.save(output_file)
        return True
    except Exception as e:
        print(f"  ❌ Error converting {txt_file.name} to DOCX: {e}")
        return False

def convert_txt_to_pdf(txt_file, output_file):
    """Convert a TXT resume to PDF format."""
    try:
        # Read TXT content
        with open(txt_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Use standard fonts that support bold (more reliable)
        # Standard fonts: Arial, Times, Courier (all support bold)
        font_name = 'Arial'
        pdf.set_font(font_name, size=10)
        
        # Process content
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                pdf.ln(5)
                continue
            
            # Check if heading
            is_heading = (
                line.isupper() and len(line) > 5 and ' ' in line
            ) or line.endswith((':', '-', '=')) or (
                len(line) < 50 and line.isupper()
            )
            
            if is_heading:
                # Heading - use bold
                pdf.set_font(font_name, 'B', 14)
                heading_text = line.rstrip(':-=').strip()
                # Replace any non-ASCII characters with ASCII equivalents for standard fonts
                try:
                    pdf.cell(0, 10, heading_text, ln=1)
                except:
                    # If encoding fails, use a simplified version
                    heading_text_ascii = heading_text.encode('ascii', 'ignore').decode('ascii')
                    pdf.cell(0, 10, heading_text_ascii, ln=1)
                pdf.set_font(font_name, '', 10)
            else:
                # Regular text
                pdf.set_font(font_name, size=10)
                
                # Always use multi_cell for better text wrapping
                # Clean the line and handle encoding
                try:
                    # Remove any problematic characters and ensure ASCII
                    line_clean = line.encode('ascii', 'ignore').decode('ascii')
                    # Remove any control characters
                    line_clean = ''.join(char for char in line_clean if ord(char) >= 32 or char in '\n\t')
                    # Ensure we have some width for multi_cell
                    if line_clean.strip():
                        pdf.multi_cell(0, 5, line_clean)
                except Exception as e:
                    # Last resort: split line into chunks
                    try:
                        # Break into smaller chunks
                        chunk_size = 100
                        for i in range(0, len(line), chunk_size):
                            chunk = line[i:i+chunk_size].encode('ascii', 'ignore').decode('ascii')
                            if chunk.strip():
                                pdf.multi_cell(0, 5, chunk)
                    except:
                        # Skip problematic lines
                        pass
        
        # Save PDF
        pdf.output(str(output_file))
        return True
    except Exception as e:
        print(f"  ❌ Error converting {txt_file.name} to PDF: {e}")
        import traceback
        traceback.print_exc()
        return False

def convert_all_test_resumes():
    """Convert all TXT resumes in data folder to PDF and DOCX."""
    data_dir = Path("data")
    
    if not data_dir.exists():
        print("❌ 'data' directory not found!")
        return
    
    # Find all resume TXT files
    resume_txt_files = sorted(data_dir.glob("resume_*.txt"))
    
    if not resume_txt_files:
        print("❌ No resume TXT files found in 'data' directory!")
        return
    
    print(f"Found {len(resume_txt_files)} resume(s) to convert\n")
    
    docx_count = 0
    pdf_count = 0
    
    for txt_file in resume_txt_files:
        base_name = txt_file.stem
        
        # Convert to DOCX
        docx_file = data_dir / f"{base_name}.docx"
        print(f"Converting {txt_file.name} → {docx_file.name}...", end=" ")
        if convert_txt_to_docx(txt_file, docx_file):
            print("✓")
            docx_count += 1
        else:
            print("✗")
        
        # Convert to PDF
        pdf_file = data_dir / f"{base_name}.pdf"
        print(f"Converting {txt_file.name} → {pdf_file.name}...", end=" ")
        if convert_txt_to_pdf(txt_file, pdf_file):
            print("✓")
            pdf_count += 1
        else:
            print("✗")
        print()
    
    print("=" * 60)
    print(f"✅ Conversion complete!")
    print(f"  - DOCX files created: {docx_count}/{len(resume_txt_files)}")
    print(f"  - PDF files created: {pdf_count}/{len(resume_txt_files)}")
    print(f"\n📁 All files saved in: {data_dir.absolute()}")

if __name__ == "__main__":
    print("Converting test resumes to PDF and DOCX formats...\n")
    convert_all_test_resumes()


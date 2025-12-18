"""
Convert text resumes to PDF and DOCX formats for testing
"""
from pathlib import Path
import os


def create_docx_from_text(text_content, output_path):
    """Create a DOCX file from text content."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Split text into lines
        lines = text_content.split('\n')
        
        for line in lines:
            if not line.strip():
                # Empty line
                doc.add_paragraph()
                continue
            
            # Check if it's a header (name or section)
            if line.isupper() or (line and line[0].isupper() and len(line) < 80 and ':' not in line):
                paragraph = doc.add_paragraph(line.strip())
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)
                
                # Make it bold if it looks like a section header
                if line.isupper() or (len(line) < 50 and not ',' in line):
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
            else:
                # Regular paragraph
                paragraph = doc.add_paragraph(line.strip())
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        doc.save(output_path)
        print(f"✓ Created DOCX: {output_path}")
        return True
    except ImportError:
        print(f"⚠ python-docx not installed. Install with: pip install python-docx")
        return False
    except Exception as e:
        print(f"✗ Error creating DOCX {output_path}: {str(e)}")
        return False


def create_pdf_from_text(text_content, output_path):
    """Create a PDF file from text content."""
    # Try fpdf2 first (lighter weight)
    try:
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        
        # Split text into lines
        lines = text_content.split('\n')
        
        for line in lines:
            if not line.strip():
                pdf.ln(5)
                continue
            
            line = line.strip()
            
            # Check if it's a header
            is_header = (line.isupper() or 
                        (line and line[0].isupper() and len(line) < 80 and 
                         ':' not in line and not ',' in line[:30]))
            
            if is_header:
                pdf.set_font("Arial", style="B", size=12)
                pdf.cell(0, 8, line, ln=1)
                pdf.set_font("Arial", size=10)
            else:
                # Handle long lines by wrapping
                pdf.multi_cell(0, 5, line)
            
            pdf.ln(3)
        
        pdf.output(str(output_path))
        print(f"✓ Created PDF: {output_path}")
        return True
        
    except ImportError:
        # Fallback to reportlab if fpdf2 not available
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT
            
            # Create PDF document
            doc = SimpleDocTemplate(str(output_path), pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor='black',
                spaceAfter=6,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor='black',
                spaceAfter=3,
                alignment=TA_LEFT,
                leading=12
            )
            
            # Split text into lines
            lines = text_content.split('\n')
            
            for line in lines:
                if not line.strip():
                    elements.append(Spacer(1, 0.1*inch))
                    continue
                
                line = line.strip()
                
                # Determine style based on line content
                if line.isupper() or (line and line[0].isupper() and len(line) < 80 and ':' not in line and not ',' in line[:30]):
                    elements.append(Paragraph(line, heading_style))
                else:
                    # Regular text - escape special characters
                    line_escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    elements.append(Paragraph(line_escaped, normal_style))
            
            # Build PDF
            doc.build(elements)
            print(f"✓ Created PDF: {output_path}")
            return True
            
        except ImportError:
            print(f"⚠ PDF libraries not installed. Install with: pip install fpdf2 OR pip install reportlab")
            return False
    except Exception as e:
        print(f"✗ Error creating PDF {output_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def convert_resumes_to_pdf_docx(data_dir="data"):
    """Convert all text resumes to PDF and DOCX formats."""
    data_path = Path(data_dir)
    
    if not data_path.exists():
        print(f"❌ Directory '{data_dir}' not found!")
        return
    
    # Find all resume text files
    resume_files = list(data_path.glob("resume_*.txt"))
    
    if not resume_files:
        print(f"⚠ No resume files found in '{data_dir}' directory.")
        print("   Run 'python generate_sample_data.py' first to create sample resumes.")
        return
    
    print(f"📄 Found {len(resume_files)} resume file(s) to convert\n")
    
    pdf_count = 0
    docx_count = 0
    
    for resume_file in resume_files:
        # Read text content
        try:
            with open(resume_file, 'r', encoding='utf-8') as f:
                text_content = f.read()
        except Exception as e:
            print(f"✗ Error reading {resume_file}: {str(e)}")
            continue
        
        # Get base filename without extension
        base_name = resume_file.stem
        
        # Create PDF
        pdf_path = data_path / f"{base_name}.pdf"
        if create_pdf_from_text(text_content, pdf_path):
            pdf_count += 1
        
        # Create DOCX
        docx_path = data_path / f"{base_name}.docx"
        if create_docx_from_text(text_content, docx_path):
            docx_count += 1
        
        print()  # Empty line between files
    
    print("="*60)
    print(f"✓ Conversion complete!")
    print(f"  - Created {pdf_count} PDF file(s)")
    print(f"  - Created {docx_count} DOCX file(s)")
    print("="*60)


if __name__ == "__main__":
    convert_resumes_to_pdf_docx()

